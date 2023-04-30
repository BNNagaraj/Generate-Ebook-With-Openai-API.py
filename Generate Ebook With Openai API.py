# Generate E-book with openai
import os
import openai
import requests
import re
from docx import Document

# Set up OpenAI API key
openai.api_key = api_key


model = "text-davinci-003"
tokens = 4000

# Function to generate content
def generate_text(prompt, model=model, tokens_per_call=tokens, num_calls=1):
    combined_text = ""
    for _ in range(num_calls):
        response = openai.Completion.create(
            engine=model,
            prompt=prompt,
            max_tokens=tokens_per_call,
            n=1,
            stop=None,
            temperature=0.7,
        )
        combined_text += response.choices[0].text.strip() + "\n\n"
    return combined_text.strip()

toc_prompt_1 = "How to be Calm and Confident in Stressful Situations"
toc_prompt_2 = f"Give engaging title to ebook on subject called '{toc_prompt_1}'"
toc_prompt_3 = generate_text(toc_prompt_2, model=model)


def main():
    # Get eBook title from user
    ebook_title = toc_prompt_3
    print(f"Generating an eBook on '{ebook_title}'...")

    # Sanitize the filename
    sanitized_filename = re.sub(r'[\\/:*?"<>|]', '', ebook_title) + ".docx"

    # Generate table of contents
    toc_prompt = f"Generate a table of contents for an eBook on '{ebook_title}':"
    toc = generate_text(toc_prompt).split('\n')
    print(toc)

    # Create a new Word document
    document = Document()

    # Add the eBook title to the document
    document.add_heading(ebook_title, level=1)

    # Generate content for each section
    for i, section in enumerate(toc):
        document.add_heading(section, level=2)
        content_prompt = f"Write a detailed content for an ebook called '{ebook_title}' on topic '{section}':"
        content = generate_text(content_prompt, model=model, tokens_per_call=tokens, num_calls=1)

        for line in content.split('\n'):
            document.add_paragraph(line.strip())

        print(content)

    # Save the generated eBook as a Word document
    document.save(sanitized_filename)
    print(f"eBook '{ebook_title}' has been saved as a Word document.")

if __name__ == "__main__":
    main()